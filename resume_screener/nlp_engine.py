"""
resume_screener/nlp_engine.py

Full NLP pipeline using NLTK + scikit-learn (TF-IDF + cosine similarity).
No external API keys required — runs 100% locally.

Pipeline:
  1. extract_text()        — PDF / DOCX → raw string
  2. preprocess()          — tokenise, lower, remove stopwords, lemmatise
  3. extract_skills()      — match against a comprehensive skill vocabulary
  4. extract_candidate_info() — name, email, phone, education heuristics
  5. tfidf_similarity()    — cosine similarity between two text blocks
  6. skill_overlap_score() — Jaccard-style skill matching
  7. analyse()             — orchestrator: returns a full result dict
"""

import re
import json
import string
import logging
import math
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Optional heavy imports (graceful fallback if not installed) ───────────────
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber not installed — PDF extraction disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX extraction disabled.")

NLTK_AVAILABLE = False
try:
    import nltk

    # ── Quietly download every resource we need. ──────────────────────────
    # We wrap each download individually so one failure never blocks the rest.
    # We use download(..., quiet=True, raise_on_error=False) — that flag
    # tells NLTK to return False instead of raising on network errors.
    _NLTK_PACKAGES = [
        'stopwords',
        'wordnet',
        'punkt',
        'punkt_tab',
        'averaged_perceptron_tagger',
        'omw-1.4',          # WordNet's multilingual extension (needed on some platforms)
    ]
    for _pkg in _NLTK_PACKAGES:
        try:
            nltk.download(_pkg, quiet=True, raise_on_error=False)
        except Exception:
            pass  # network down, SSL error, KeyboardInterrupt — carry on

    # ── Now try to import the modules we actually use. ─────────────────────
    # If a resource is still missing (e.g. offline machine), the functions
    # that use NLTK will fall back to basic string splitting — see preprocess().
    try:
        from nltk.corpus import stopwords as _sw
        from nltk.stem import WordNetLemmatizer as _WNL
        from nltk.tokenize import word_tokenize as _wt
        # Quick smoke-test — this raises LookupError if data is absent
        _sw.words('english')
        stopwords      = _sw
        WordNetLemmatizer = _WNL
        word_tokenize  = _wt
        NLTK_AVAILABLE = True
    except Exception as _e:
        logger.warning("NLTK data not fully available — using basic tokenisation. (%s)", _e)
        NLTK_AVAILABLE = False

except ImportError:
    logger.warning("NLTK not installed — using basic tokenisation.")

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False
    logger.warning("scikit-learn not installed — TF-IDF scoring disabled.")


# ─────────────────────────────────────────────────────────────────────────────
# COMPREHENSIVE SKILL VOCABULARY
# ─────────────────────────────────────────────────────────────────────────────

SKILL_VOCAB = {
    # Programming languages
    'python', 'java', 'javascript', 'typescript', 'c', 'c++', 'c#', 'ruby',
    'php', 'swift', 'kotlin', 'go', 'golang', 'rust', 'scala', 'r',
    'matlab', 'perl', 'bash', 'shell', 'powershell', 'lua', 'dart', 'elixir',
    'haskell', 'clojure', 'f#', 'objective-c', 'assembly', 'cobol', 'fortran',

    # Web frontend
    'html', 'css', 'sass', 'scss', 'less', 'react', 'reactjs', 'angular',
    'angularjs', 'vue', 'vuejs', 'svelte', 'jquery', 'bootstrap', 'tailwind',
    'tailwindcss', 'material-ui', 'chakra-ui', 'next.js', 'nextjs', 'nuxt',
    'gatsby', 'webpack', 'vite', 'babel', 'eslint', 'redux', 'mobx',
    'graphql', 'apollo', 'websocket', 'pwa', 'web components',

    # Web backend
    'django', 'flask', 'fastapi', 'express', 'expressjs', 'node.js', 'nodejs',
    'spring', 'spring boot', 'hibernate', 'laravel', 'rails', 'ruby on rails',
    'asp.net', '.net', 'dotnet', 'gin', 'echo', 'fiber', 'phoenix',
    'actix', 'nestjs', 'strapi', 'hasura',

    # Databases
    'sql', 'mysql', 'postgresql', 'postgres', 'sqlite', 'oracle', 'mssql',
    'sql server', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
    'dynamodb', 'firebase', 'supabase', 'neo4j', 'couchdb', 'influxdb',
    'timescaledb', 'mariadb', 'cockroachdb', 'prisma', 'sqlalchemy',

    # Cloud & DevOps
    'aws', 'amazon web services', 'azure', 'gcp', 'google cloud',
    'docker', 'kubernetes', 'k8s', 'terraform', 'ansible', 'puppet',
    'chef', 'jenkins', 'github actions', 'gitlab ci', 'circleci',
    'travis ci', 'argocd', 'helm', 'istio', 'nginx', 'apache',
    'linux', 'ubuntu', 'centos', 'debian', 'ci/cd', 'devops',
    'sre', 'site reliability', 'prometheus', 'grafana', 'datadog',
    'splunk', 'elk stack', 'logstash', 'kibana',

    # Data Science / ML / AI
    'machine learning', 'deep learning', 'neural networks', 'nlp',
    'natural language processing', 'computer vision', 'data science',
    'data analysis', 'data engineering', 'feature engineering',
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'sklearn',
    'pandas', 'numpy', 'scipy', 'matplotlib', 'seaborn', 'plotly',
    'tableau', 'power bi', 'looker', 'spark', 'apache spark', 'hadoop',
    'kafka', 'airflow', 'dbt', 'mlflow', 'hugging face', 'transformers',
    'bert', 'gpt', 'llm', 'langchain', 'openai', 'computer vision',
    'opencv', 'yolo', 'statistics', 'a/b testing', 'sql', 'etl',

    # Mobile
    'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic',
    'swift', 'objective-c', 'kotlin', 'dart', 'xcode', 'android studio',

    # Security
    'cybersecurity', 'penetration testing', 'ethical hacking', 'owasp',
    'ssl', 'tls', 'oauth', 'jwt', 'encryption', 'pki', 'soc', 'siem',
    'vulnerability assessment', 'firewalls', 'ids', 'ips', 'zero trust',

    # Tools & Practices
    'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence',
    'slack', 'figma', 'sketch', 'photoshop', 'illustrator',
    'agile', 'scrum', 'kanban', 'waterfall', 'tdd', 'bdd',
    'test driven development', 'unit testing', 'integration testing',
    'code review', 'pair programming', 'microservices', 'monolith',
    'rest', 'restful', 'api', 'soap', 'grpc', 'message queue',
    'rabbitmq', 'sqs', 'pub/sub', 'event driven', 'cqrs', 'ddd',

    # Soft skills
    'leadership', 'communication', 'teamwork', 'problem solving',
    'critical thinking', 'project management', 'time management',
    'collaboration', 'mentoring', 'presentation', 'documentation',
    'analytical', 'research', 'planning', 'stakeholder management',
}


# ─────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_path: str) -> str:
    """
    Extract raw text from a PDF or DOCX file.
    Returns empty string if extraction fails.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == '.pdf':
        return _extract_pdf(path)
    elif suffix == '.docx':
        return _extract_docx(path)
    else:
        logger.error("Unsupported file type: %s", suffix)
        return ''


def _extract_pdf(path: Path) -> str:
    if not PDF_AVAILABLE:
        return ''
    pages = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    pages.append(text)
                else:
                    logger.debug("PDF page %d: no text extracted (may be image-based)", i + 1)
        result = '\n\n'.join(pages).strip()
        logger.info("PDF extracted: %d chars from %s", len(result), path.name)
        return result
    except Exception as e:
        logger.error("PDF extraction failed for %s: %s", path.name, e)
        return ''


def _extract_docx(path: Path) -> str:
    if not DOCX_AVAILABLE:
        return ''
    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        result = '\n'.join(paragraphs).strip()
        logger.info("DOCX extracted: %d chars from %s", len(result), path.name)
        return result
    except Exception as e:
        logger.error("DOCX extraction failed for %s: %s", path.name, e)
        return ''


# ─────────────────────────────────────────────────────────────────────────────
# TEXT PREPROCESSING
# ─────────────────────────────────────────────────────────────────────────────

def preprocess(text: str) -> str:
    """
    Lowercase, remove punctuation, remove stopwords, lemmatise.
    Falls back to simple cleaning if NLTK is unavailable.
    """
    if not text:
        return ''

    text = text.lower()
    # Remove URLs
    text = re.sub(r'http\S+|www\.\S+', ' ', text)
    # Remove email addresses
    text = re.sub(r'\S+@\S+', ' ', text)
    # Remove phone numbers
    text = re.sub(r'[\+\(]?[0-9][0-9\s\-\(\)]{7,}[0-9]', ' ', text)
    # Remove special chars (keep hyphens inside words like "full-stack")
    text = re.sub(r'[^\w\s\-]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    if not NLTK_AVAILABLE:
        return text

    try:
        tokens = word_tokenize(text)
        stop_words = set(stopwords.words('english'))
        lemmatiser = WordNetLemmatizer()
        tokens = [
            lemmatiser.lemmatize(t)
            for t in tokens
            if t not in stop_words and len(t) > 1
        ]
        return ' '.join(tokens)
    except Exception as e:
        logger.warning("NLTK preprocessing failed, using raw text: %s", e)
        return text


# ─────────────────────────────────────────────────────────────────────────────
# SKILL EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_skills(text: str) -> list[str]:
    """
    Find all skills from SKILL_VOCAB present in the text.
    Uses word-boundary matching so 'r' doesn't match 'react'.
    Multi-word skills (e.g. 'machine learning') are matched as phrases.
    """
    if not text:
        return []

    text_lower = text.lower()
    found = set()

    for skill in SKILL_VOCAB:
        # Escape dots and pluses in skill names (e.g. 'c++', 'asp.net')
        escaped = re.escape(skill)
        # For single-word skills use \b; multi-word skills use lookaround spaces
        if ' ' in skill:
            pattern = rf'(?<!\w){escaped}(?!\w)'
        else:
            pattern = rf'\b{escaped}\b'

        if re.search(pattern, text_lower):
            found.add(skill)

    return sorted(found)


# ─────────────────────────────────────────────────────────────────────────────
# CANDIDATE INFO EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_candidate_info(text: str) -> dict:
    """
    Heuristically extract name, email, phone, years of experience.
    """
    info = {
        'name': 'Unknown',
        'email': '',
        'phone': '',
        'experience_years': None,
    }

    if not text:
        return info

    # Email
    emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
    if emails:
        info['email'] = emails[0]

    # Phone
    phones = re.findall(
        r'(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text
    )
    if phones:
        info['phone'] = phones[0]

    # Name: first non-empty line that looks like a name (2–4 capitalised words)
    for line in text.split('\n')[:15]:
        line = line.strip()
        if not line or '@' in line or re.search(r'\d{4}', line):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
            # Skip lines that are obviously section headers
            if not any(kw in line.lower() for kw in
                       ('experience', 'education', 'skill', 'summary', 'objective', 'profile')):
                info['name'] = line
                break

    # Years of experience
    exp_patterns = [
        r'(\d+)\+?\s*years?\s+of\s+experience',
        r'(\d+)\+?\s*years?\s+experience',
        r'experience\s+of\s+(\d+)\+?\s*years?',
    ]
    for pattern in exp_patterns:
        m = re.search(pattern, text.lower())
        if m:
            info['experience_years'] = int(m.group(1))
            break

    return info


# ─────────────────────────────────────────────────────────────────────────────
# TF-IDF COSINE SIMILARITY
# ─────────────────────────────────────────────────────────────────────────────

def tfidf_similarity(text_a: str, text_b: str) -> float:
    """
    Compute cosine similarity between two preprocessed texts using TF-IDF.
    Returns a float in [0.0, 1.0].
    """
    if not SKLEARN_AVAILABLE:
        return _fallback_similarity(text_a, text_b)

    if not text_a.strip() or not text_b.strip():
        return 0.0

    try:
        vectorizer = TfidfVectorizer(
            ngram_range=(1, 2),       # unigrams + bigrams
            min_df=1,
            sublinear_tf=True,        # log-scale TF dampening
            max_features=10_000,
        )
        matrix = vectorizer.fit_transform([text_a, text_b])
        sim = cosine_similarity(matrix[0:1], matrix[1:2])[0][0]
        return float(sim)
    except Exception as e:
        logger.warning("TF-IDF similarity failed: %s", e)
        return _fallback_similarity(text_a, text_b)


def _fallback_similarity(text_a: str, text_b: str) -> float:
    """
    Pure-Python fallback: Jaccard similarity on word sets.
    Less accurate than TF-IDF but requires no dependencies.
    """
    set_a = set(text_a.lower().split())
    set_b = set(text_b.lower().split())
    if not set_a or not set_b:
        return 0.0
    intersection = set_a & set_b
    union = set_a | set_b
    return len(intersection) / len(union)


# ─────────────────────────────────────────────────────────────────────────────
# SKILL OVERLAP SCORE
# ─────────────────────────────────────────────────────────────────────────────

def skill_overlap_score(candidate_skills: list, required_skills: list) -> tuple[float, list, list]:
    """
    Returns (score_0_to_100, matched_skills, missing_skills).
    Score = matched / required  (penalises missing required skills).
    """
    if not required_skills:
        return (50.0, [], [])  # No requirements = neutral score

    candidate_set = {s.lower() for s in candidate_skills}
    required_set  = {s.lower() for s in required_skills}

    matched = sorted(candidate_set & required_set)
    missing = sorted(required_set - candidate_set)

    score = (len(matched) / len(required_set)) * 100.0
    return (round(score, 2), matched, missing)


# ─────────────────────────────────────────────────────────────────────────────
# NLP SUMMARY GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_summary(
    match_score: float,
    skill_score: float,
    tfidf_score_pct: float,
    matched: list,
    missing: list,
    candidate_name: str,
    job_title: str,
) -> str:
    """
    Generate a plain-English paragraph summarising the NLP analysis results.
    """
    label = (
        'an excellent' if match_score >= 80 else
        'a good'       if match_score >= 65 else
        'a moderate'   if match_score >= 45 else
        'a low'
    )

    matched_str = ', '.join(matched[:6]) if matched else 'none detected'
    missing_str = ', '.join(missing[:5]) if missing else 'none'

    lines = [
        f"{candidate_name} is {label} match for the {job_title} role "
        f"with an overall AI score of {match_score:.1f}%.",
        "",
        f"Skill coverage: {skill_score:.1f}% of required skills detected. "
        f"Matched skills include: {matched_str}.",
    ]

    if missing:
        lines.append(
            f"Skills to develop or highlight: {missing_str}."
        )
    else:
        lines.append("All required skills were found in the resume — outstanding coverage.")

    lines += [
        "",
        f"Semantic relevance (TF-IDF cosine similarity): {tfidf_score_pct:.1f}% — "
        + ("strong contextual alignment with the job description." if tfidf_score_pct >= 50
           else "moderate contextual alignment with the job description.")
    ]

    return '\n'.join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ANALYSIS ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────

def analyse(resume_file_path: str, job_posting) -> dict:
    """
    Full NLP pipeline. Call this from the Django view.

    Args:
        resume_file_path : absolute path to the uploaded PDF/DOCX
        job_posting      : JobPosting model instance

    Returns a dict:
      {
        'extracted_text'   : str,
        'candidate_info'   : dict,
        'candidate_skills' : list[str],
        'matched_skills'   : list[str],
        'missing_skills'   : list[str],
        'match_score'      : float,    # 0-100 combined score
        'skill_score'      : float,    # 0-100 skill overlap
        'tfidf_score'      : float,    # 0-100 semantic score
        'nlp_summary'      : str,
        'error'            : str | None,
      }
    """
    result = {
        'extracted_text':   '',
        'candidate_info':   {},
        'candidate_skills': [],
        'matched_skills':   [],
        'missing_skills':   [],
        'match_score':      0.0,
        'skill_score':      0.0,
        'tfidf_score':      0.0,
        'nlp_summary':      '',
        'error':            None,
    }

    # ── Step 1: Extract text ──────────────────────────────────────────────
    raw_text = extract_text(resume_file_path)
    if not raw_text:
        result['error'] = (
            'Could not extract text from the uploaded file. '
            'Ensure the PDF is text-based (not a scanned image) or use a DOCX file.'
        )
        result['nlp_summary'] = result['error']
        return result

    result['extracted_text'] = raw_text

    # ── Step 2: Extract candidate info ────────────────────────────────────
    result['candidate_info'] = extract_candidate_info(raw_text)

    # ── Step 3: Extract skills from resume ───────────────────────────────
    candidate_skills = extract_skills(raw_text)
    result['candidate_skills'] = candidate_skills

    # ── Step 4: Get required skills from job posting ──────────────────────
    required_skills = job_posting.get_required_skills_list()

    # ── Step 5: Skill overlap score ───────────────────────────────────────
    skill_score, matched, missing = skill_overlap_score(candidate_skills, required_skills)
    result['skill_score']    = skill_score
    result['matched_skills'] = matched
    result['missing_skills'] = missing

    # ── Step 6: TF-IDF semantic similarity ───────────────────────────────
    processed_resume = preprocess(raw_text)
    processed_jd     = preprocess(job_posting.description + ' ' + job_posting.required_skills)

    raw_tfidf   = tfidf_similarity(processed_resume, processed_jd)
    tfidf_score = round(raw_tfidf * 100, 2)
    result['tfidf_score'] = tfidf_score

    # ── Step 7: Combined score (weighted average) ─────────────────────────
    # 60% skill overlap (most directly relevant)
    # 40% TF-IDF semantic match (context and terminology)
    match_score = round((skill_score * 0.60) + (tfidf_score * 0.40), 2)
    match_score = max(0.0, min(100.0, match_score))
    result['match_score'] = match_score

    # ── Step 8: Generate NLP summary ─────────────────────────────────────
    candidate_name = result['candidate_info'].get('name', 'The candidate')
    result['nlp_summary'] = generate_summary(
        match_score=match_score,
        skill_score=skill_score,
        tfidf_score_pct=tfidf_score,
        matched=matched,
        missing=missing,
        candidate_name=candidate_name,
        job_title=job_posting.title,
    )

    logger.info(
        "NLP analysis complete | Job: '%s' | Score: %.1f%% | Skills: %d/%d",
        job_posting.title,
        match_score,
        len(matched),
        len(required_skills),
    )

    return result
